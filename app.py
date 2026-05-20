"""
TK Listing Agent — TikTok Shop Listing Generator.
Gradio UI connecting Apify, Bright Data, DeepSeek, Seedream 5.0, Qianwen.
Run: python app.py
"""

import os
import io
import json
import zipfile
import tempfile

import gradio as gr

from apify_fetcher import search_products_apify
from data_fetcher import search_posts_by_keyword, get_comments, analyze_comments
from text_generator import (
    generate_titles,
    generate_description,
    generate_hashtags,
    extract_keywords,
)
from image_generator_v2 import (
    generate_product_photos,
    generate_detail_module_b1,
    generate_detail_module_b2,
    generate_single_feature_deep,
    generate_comparison_card,
    generate_detail_module_b3,
    generate_detail_module_b4,
    generate_detail_module_b5,
    generate_detail_module_b11,
    generate_detail_module_b12,
    generate_detail_module_b13,
    stitch_detail_long_image,
)

# ============================================================
# Local cache — persists across page refreshes
# ============================================================
CACHE_DIR = os.path.join(os.path.dirname(__file__), ".cache")
os.makedirs(CACHE_DIR, exist_ok=True)

def _cache_put(key: str, value):
    with open(os.path.join(CACHE_DIR, f"{key}.json"), "w", encoding="utf-8") as f:
        json.dump(value, f, ensure_ascii=False)

def _cache_get(key: str):
    p = os.path.join(CACHE_DIR, f"{key}.json")
    if os.path.exists(p):
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

def _cache_put_img(key: str, img):
    img.save(os.path.join(CACHE_DIR, f"{key}.jpg"), format="JPEG", quality=95)


# ============================================================
# Global generated-images tracker (for Tab7 dropdowns)
# ============================================================

generated_images: list[dict] = []  # [{"path": str, "type": str, "label": str}]

# Dropdown slot names (for generating choices)
DROPDOWN_SLOTS = ["main_image", "detail_1", "detail_2", "scene_image", "model_image"]

DEFAULT_CHOICE = ("无（使用占位）", "")


def _build_dropdown_choices() -> list[tuple[str, str]]:
    """Build (label, value) choices list from generated_images."""
    choices = [DEFAULT_CHOICE]
    for img in generated_images:
        label = f"{img['type']} - {img['label']}"
        choices.append((label, img["path"]))
    return choices


def refresh_image_dropdowns():
    """Return 5 gr.update() for the 5 image-slot dropdowns."""
    choices = _build_dropdown_choices()
    return tuple(gr.update(choices=choices, value="") for _ in range(5))


# ============================================================
# Callback functions
# ============================================================

def _search_posts_with_timeout(keyword: str, limit: int = 20, timeout: int = 600) -> dict:
    """Run Bright Data posts search with a shared list for partial results."""
    import threading

    result: list[dict] = []
    exc: Exception | None = None
    done = threading.Event()

    def _run():
        nonlocal result, exc
        try:
            result = search_posts_by_keyword(keyword, limit=limit)
        except Exception as e:
            exc = e
        finally:
            done.set()

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    got_result = done.wait(timeout=timeout)
    if got_result:
        if exc:
            raise exc
        return {"posts": result, "timed_out": False}
    else:
        print(f"[posts] timeout after {timeout}s, returning partial data")
        return {"posts": result if result else [], "timed_out": True}


def on_search(state: dict, keyword: str, country: str, progress=gr.Progress()):
    """Two-step search: Apify immediately → Tab2, then Bright Data background → Tab3.
    Uses generator to yield intermediate updates so Tab2 appears right after Apify.
    Returns: yields (state, status_md, df_rows, posts_md)
    """
    state = dict(state)
    status = ""

    if not keyword.strip():
        yield state, "  请输入搜索关键词。", [], ""
        return

    # ---- Progress tracking ----
    progress(0.1, desc="正在通过 Apify 搜索商品……")

    # ---- Step 1: Apify products (runs immediately, updates Tab2) ----
    products = []
    try:
        products = search_products_apify(keyword, country, limit=10)
        state["products"] = products
    except Exception as e:
        state["products"] = []
        status += f"  Apify 搜索失败：{e}\n"

    n_products = len(products)
    is_mock = False

    # Fallback: if Apify returns empty, retry with mock data
    if n_products == 0:
        print(f"[on_search] Apify returned 0 results for '{keyword}', falling back to mock data")
        products = search_products_apify(keyword, country, limit=10, use_mock=True)
        state["products"] = products
        n_products = len(products)
        is_mock = True

    # Build Tab2 dataframe
    df_rows = [[p.get("title", ""), p.get("price", ""), p.get("url", "")] for p in products]

    # ---- Yield intermediate: Tab2 visible with products, Tab3 shows pending ----
    if n_products > 0:
        if is_mock:
            status += (
                '  <span style="color: #DAA520; font-weight: bold;">'
                "⚠️ 已加载 5 条手串商品数据（模拟数据）"
                "</span>\n"
            )
        else:
            status += f"  ✅ 商品搜索完成！找到 {n_products} 个商品。\n"
    else:
        status += (
            '  <span style="color: #DAA520; font-weight: bold;">'
            "⚠️ 未找到商品数据，请尝试更通用的关键词（如 healing bracelet、gemstone bracelet）"
            "</span>\n"
        )
    status += "  🕐 正在后台获取帖子趋势……\n"
    yield state, status, df_rows, "*正在搜索帖子趋势……*\n\n> 请稍候，商品列表已就绪，可切换到 Tab 2 查看。"

    # ---- Step 2: Bright Data posts (background, up to 5 min) ----
    progress(0.4, desc="正在通过 Bright Data 搜索热门帖子……")
    n_posts = 0
    try:
        posts_result = _search_posts_with_timeout(keyword, limit=20, timeout=300)
        posts = posts_result["posts"]
        state["posts"] = posts
        n_posts = len(posts)
    except Exception as e:
        posts = []
        state["posts"] = []
        status += f"  帖子搜索失败：{e}\n"

    # Build Tab3 posts markdown
    posts = state.get("posts") or []
    if posts:
        lines = [f"###  找到 {len(posts)} 条帖子\n"]
        for i, p in enumerate(posts[:10]):
            title = p.get("title", "(untitled)")[:100]
            tags = " ".join(p.get("tags", [])[:6])
            plays = p.get("play_count", 0)
            lines.append(f"**{i+1}.** {title}")
            lines.append(f"> 播放量：{plays} | 标签：{tags}")
            lines.append("")
        posts_md = "\n".join(lines)
    else:
        posts_md = "*暂无热门帖子。*"

    # ---- Final yield: everything done ----
    progress(1.0, desc="搜索完成！")
    if n_posts > 0:
        status += f"  ✅ 帖子搜索完成！找到 {n_posts} 条热门帖子。\n"
    elif "失败" in status:
        pass
    else:
        status += '  <span style="color: #DAA520;">⚠️ 帖子搜索无结果，文案仍可基于商品数据生成。</span>\n'
    status += "\n  ✅ 全部数据搜索完成！请切换到 **Tab 2 或 3** 查看结果。"

    # Cache search results
    import json as _json
    _cache_put("products", state.get("products", []))
    _cache_put("posts", state.get("posts", []))

    yield state, status, df_rows, posts_md


def on_generate_copy(state: dict, brand: str, main_kw: str, category_kw: str,
                     selling_pts: str, crowd_kw: str):
    """Generate titles, description, hashtags from search results + inputs.
    Returns: (state, status, seo_title, conv_title, brand_title, sp_md, qa_md, cautions_md, hashtags_str)
    """
    state = dict(state)
    status = ""

    # Extract hot keywords from collected data
    status += "  正在提取关键词……\n"
    try:
        texts = []
        for p in (state.get("products") or []):
            if p.get("title"):
                texts.append(p["title"])
        for p in (state.get("posts") or []):
            if p.get("title"):
                texts.append(p["title"])
        if not texts:
            texts = [main_kw, selling_pts]
        hot_kw = extract_keywords(texts)
    except Exception as e:
        hot_kw = main_kw
        status += f"  关键词提取失败：{e}\n"

    # Titles
    status += "  正在生成标题……\n"
    seo_title = conv_title = brand_title = ""
    try:
        titles = generate_titles(
            brand=brand, main_keyword=main_kw,
            hot_keywords=hot_kw, crowd_keywords=crowd_kw,
            selling_points=selling_pts,
        )
        state["titles"] = titles
        seo_title = titles[0] if len(titles) > 0 else ""
        conv_title = titles[1] if len(titles) > 1 else ""
        brand_title = titles[2] if len(titles) > 2 else ""
    except Exception as e:
        status += f"  标题生成失败：{e}\n"

    # Description
    status += "  正在生成描述……\n"
    sp_md = qa_md = cautions_md = ""
    try:
        desc = generate_description(
            brand=brand, main_keyword=main_kw,
            selling_points=selling_pts,
        )
        state["description"] = desc
        sp = desc.get("selling_points", [])
        qa = desc.get("qa", [])
        ca = desc.get("cautions", [])
        sp_md = "\n".join(f"- {s.replace(chr(10), ' | ')}" for s in sp) if sp else "*暂无卖点*"
        qa_md = "\n".join(f"- {q}" for q in qa) if qa else "*暂无问答*"
        cautions_md = "\n".join(f"- {c}" for c in ca) if ca else "*暂无注意事项*"
    except Exception as e:
        status += f"  描述生成失败：{e}\n"

    # Hashtags
    status += "  正在生成话题标签……\n"
    hashtags_str = ""
    try:
        tags = generate_hashtags(
            main_keyword=main_kw,
            category_keyword=category_kw,
            scene_keywords="gym, home workout, daily wear",
            crowd_keywords=crowd_kw,
        )
        state["hashtags"] = tags
        hashtags_str = " ".join(tags)
    except Exception as e:
        status += f"  话题标签生成失败：{e}\n"

    status += "\n  文案生成完成！"

    # Cache copy results
    _cache_put("titles", state.get("titles", []))
    _cache_put("description", state.get("description", {}))
    _cache_put("hashtags", state.get("hashtags", []))
    return state, status, seo_title, conv_title, brand_title, sp_md, qa_md, cautions_md, hashtags_str


def on_generate_product(state: dict, ref_upload, product_keyword: str, brand: str):
    """Generate 9 product photos via Seedream 5.0. Saves to state.
    Returns: (state, status, gallery_images)
    """
    state = dict(state)

    if ref_upload is None:
        return state, "  请先上传参考图片。", []

    img_path = None
    if isinstance(ref_upload, str):
        img_path = ref_upload
    elif hasattr(ref_upload, 'name'):
        img_path = ref_upload.name

    if not img_path or not os.path.exists(img_path):
        return state, "  无法读取上传的图片，请重试。", []

    try:
        photos = generate_product_photos(
            reference_image_path=img_path,
            product_name=product_keyword or "Product",
            brand=brand or "",
            num_images=9,
        )
        if not photos:
            return state, "  生成失败，请检查 Seedream API Key 或更换参考图片后重试。", []

        # Save to state and to temp files
        state["product_photos"] = photos

        # Save to generated_images for Tab7 picker
        import tempfile
        gallery = []
        for idx, p in enumerate(photos):
            if p["image"]:
                tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False, prefix="tk_product_")
                p["image"].save(tmp.name, format="JPEG", quality=95)
                generated_images.append({
                    "path": tmp.name,
                    "type": f"商品图{idx+1}",
                    "label": f"Tab5 商品图{idx+1}",
                })
                gallery.append(p["image"])

        # Cache product photos to disk
        for idx, p in enumerate(photos):
            if p["image"]:
                _cache_put_img(f"product_{idx+1:02d}", p["image"])

        return (
            state,
            f"  ✅ 9张商品图生成完成！已加入 Tab7 图片库。现在可切换到 Tab7 生成详情长图。",
            gallery,
        )
    except Exception as e:
        return state, f"  产品图片生成失败：{e}", []


def on_clear_product(state: dict):
    """Clear product preview.
    Returns: (state, status, [])
    """
    state = dict(state)
    state.pop("product_photos", None)
    ptype = state.pop("product_image_type", None)

    # Remove matching entries from generated_images
    global generated_images
    if ptype:
        generated_images = [
            img for img in generated_images
            if not (img.get("type") == ptype and "Tab5" in img.get("label", ""))
        ]

    return state, "  已清除。上传参考图，点击生成9张商品图。", []


def on_generate_detail_v2(
    state: dict,
    brand: str, main_kw: str, material: str,
    tagline: str, material_desc: str,
    scene1: str, scene2: str, scene3: str,
):
    """Generate full 13-section detail page via Seedream product photos + Qianwen modules.
    Uses product_photos already in state (from Tab5).
    Returns: (state, status_md, pil_image)
    """
    state = dict(state)
    status = ""

    # ---- Check prerequisites ----
    photos = state.get("product_photos") or []
    if not photos:
        return state, "  ❌ 请先在 Tab 5 生成商品图（9张）。", None

    desc = state.get("description") or {}
    sp_raw = desc.get("selling_points", [])
    if not sp_raw:
        return state, "  ❌ 请先在 Tab 4 生成文案。", None

    # Parse selling points: title + description pairs
    sp_list = []
    for s in sp_raw:
        lines = s.replace(chr(10), "\n").split("\n")
        if lines:
            sp_list.append(lines[0].strip() if lines[0] else "")
    while len(sp_list) < 8:
        sp_list.append("Premium Feature")

    # Get reference images from photos (by position index)
    imgs = [p["image"] for p in photos if p["image"]]

    def _get_img(idx):
        if idx < len(imgs):
            return imgs[idx]
        return imgs[0] if imgs else None

    # Save to temp files for Qianwen API
    import tempfile
    def _save_temp(img, prefix):
        if img is None:
            return None
        t = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False, prefix=prefix)
        img.save(t.name, format="JPEG", quality=95)
        t.close()
        return t

    refs = [_save_temp(imgs[i], f"tk_ref{i}_") for i in range(min(len(imgs), 6))]
    while len(refs) < 6:
        refs.append(refs[0] if refs else _save_temp(imgs[0], "tk_ref_fb_"))

    def _r(i):
        return refs[i].name if i < len(refs) and refs[i] else refs[0].name

    # ---- Build pairs ----
    sp_pairs = []
    for i in range(0, min(len(sp_list), 8), 2):
        t = sp_list[i] if i < len(sp_list) else ""
        d = sp_list[i+1] if i+1 < len(sp_list) else ""
        sp_pairs.append((t, d))

    scenes_text = [s for s in [scene1, scene2, scene3] if s.strip()]
    if not scenes_text:
        scenes_text = ["Daily Use", "Office Style", "Weekend Casual"]

    specs_d = {}
    if material.strip():
        specs_d["Material"] = material.strip()
    specs_d["Weight"] = "Lightweight"

    # ---- Generate all 13 sections ----
    status += "  🎨 正在生成13段详情模块……\n"

    b1 = generate_detail_module_b1(white_bg_image_path=_r(0), brand=brand, product_name=main_kw, tagline=tagline)

    b2_feats = [{"title": sp_pairs[i][0][:30], "desc": sp_pairs[i][1][:60]} for i in range(4)]
    b2 = generate_detail_module_b2(reference_image_path=_r(0), features=b2_feats)

    ref_idx = [1, 2, 3, 4]
    b3 = generate_single_feature_deep(_r(ref_idx[0]), sp_pairs[0][0], sp_pairs[0][1], 1) if len(sp_pairs) > 0 else None
    b4 = generate_single_feature_deep(_r(ref_idx[1]), sp_pairs[1][0], sp_pairs[1][1], 2) if len(sp_pairs) > 1 else None
    b5 = generate_single_feature_deep(_r(ref_idx[2]), sp_pairs[2][0], sp_pairs[2][1], 3) if len(sp_pairs) > 2 else None
    b6 = generate_single_feature_deep(_r(ref_idx[3]), sp_pairs[3][0], sp_pairs[3][1], 4) if len(sp_pairs) > 3 else None

    b7 = generate_comparison_card(reference_image_path=_r(1), product_name=main_kw)
    b8 = generate_detail_module_b3(reference_image_path=_r(1), material_desc=material_desc, specs_d=specs_d)
    b9 = generate_detail_module_b4(reference_image_path=_r(3), scene_descriptions=scenes_text)
    b10 = generate_detail_module_b5(reference_image_path=_r(4), brand=brand)
    b11 = generate_detail_module_b11(specs=specs_d)
    b12 = generate_detail_module_b12(brand=brand, product_name=main_kw)
    b13 = generate_detail_module_b13(brand=brand)

    detail_long = stitch_detail_long_image(
        b1=b1, b2=b2, b3=b3, b4=b4, b5=b5, b6=b6,
        b7=b7, b8=b8, b9=b9, b10=b10, b11=b11,
        b12=b12, b13=b13,
    )

    # Cleanup temp files
    for t in refs:
        try: os.unlink(t.name)
        except: pass

    if detail_long:
        state["detail_image"] = detail_long
        # Cache detail long image
        _cache_put_img("detail_long", detail_long)
        ok_count = sum(1 for x in [b1,b2,b3,b4,b5,b6,b7,b8,b9,b10,b11,b12,b13] if x is not None)
        return state, f"  ✅ 详情长图生成完成！（{ok_count}/13段成功）", detail_long
    else:
        return state, "  ❌ 详情图生成失败，请重试。", None


def on_export(state: dict):
    """Export ALL content as a complete ZIP: Word, Excel, images, data.
    Reads from local cache first, state as fallback.
    Returns: (status_md, zip_path)
    """
    try:
        # Read from cache (survives page refresh)
        titles = _cache_get("titles") or state.get("titles") or []
        desc = _cache_get("description") or state.get("description") or {}
        hashtags = _cache_get("hashtags") or state.get("hashtags") or []
        products = _cache_get("products") or state.get("products") or []
        posts = _cache_get("posts") or state.get("posts") or []

        has_titles = bool(titles)
        has_desc = bool(desc)
        if not (has_titles or has_desc):
            return "  尚无内容可导出，请先在 Tab 4 生成文案。", None

        sp_list = desc.get("selling_points", [])
        qa_list = desc.get("qa", [])
        ca_list = desc.get("cautions", [])

        # Build keywords from products + posts
        from collections import Counter
        import re
        all_words = []
        for src in products + posts:
            text = src.get("title", "") + " " + " ".join(src.get("tags", []))
            all_words.extend(re.findall(r"[a-zA-Z]{3,}", str(text).lower()))
        stopwords = {"the","a","an","is","are","was","were","be","been","being","have","has","had",
                     "do","does","did","will","would","could","should","may","might","can","shall",
                     "you","your","they","them","their","we","us","our","this","that","these","those",
                     "it","its","and","but","or","nor","not","so","as","at","by","for","from","in",
                     "into","of","on","onto","to","with","about","up","out","if","very","just","all",
                     "also","more","some","get","got","new"}
        filtered = [w for w in all_words if w not in stopwords]
        word_counts = Counter(filtered).most_common(30)

        tmp = tempfile.mkdtemp()
        zip_path = os.path.join(tmp, "tk_listing_export.zip")

        # ======== Generate DOCX inline ========
        def _make_docx():
            from docx import Document as DocxDoc
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = DocxDoc()
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)

            doc.add_heading('TK Shop Listing Report', level=0)
            doc.add_paragraph('')

            doc.add_heading('Product Titles', level=1)
            labels = ['SEO Title', 'Conversion Title', 'Brand Title']
            for i, t in enumerate(titles[:3]):
                p = doc.add_paragraph()
                p.add_run(f'{labels[i]}: ').bold = True
                p.add_run(t)

            doc.add_heading('Selling Points', level=1)
            for sp in sp_list:
                doc.add_paragraph(sp.replace(chr(10), ' — '), style='List Bullet')

            doc.add_heading('Q&A', level=1)
            for qa in qa_list:
                doc.add_paragraph(qa.replace(chr(10), ' | '))

            doc.add_heading('Cautions', level=1)
            for ca in ca_list:
                doc.add_paragraph(ca.replace(chr(10), ' '), style='List Bullet')

            doc.add_heading('Hashtags', level=1)
            doc.add_paragraph(' '.join(hashtags))

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        # ======== Generate XLSX inline ========
        def _make_xlsx():
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            wb = Workbook()
            header_font = Font(bold=True, color='FFFFFF')
            header_fill = PatternFill('solid', start_color='2C2C2C')
            header_align = Alignment(horizontal='center')
            body_font = Font(name='Arial', size=11)

            # Sheet1: Keywords
            ws1 = wb.active
            ws1.title = 'Keywords'
            for c, h in enumerate(['Keyword', 'Count', 'Source'], 1):
                cell = ws1.cell(row=1, column=c, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align
            for r, (word, count) in enumerate(word_counts, 2):
                ws1.cell(row=r, column=1, value=word).font = body_font
                ws1.cell(row=r, column=2, value=count).font = body_font
                ws1.cell(row=r, column=3, value='Data Fetcher').font = body_font
            ws1.column_dimensions['A'].width = 24
            ws1.column_dimensions['B'].width = 10
            ws1.column_dimensions['C'].width = 16

            # Sheet2: Hashtags
            ws2 = wb.create_sheet('Hashtags')
            for c, h in enumerate(['Hashtag', 'Platform'], 1):
                cell = ws2.cell(row=1, column=c, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align
            for r, tag in enumerate(hashtags, 2):
                ws2.cell(row=r, column=1, value=tag).font = body_font
                ws2.cell(row=r, column=2, value='TikTok').font = body_font
            ws2.column_dimensions['A'].width = 28
            ws2.column_dimensions['B'].width = 14

            # Sheet3: Products
            ws3 = wb.create_sheet('Products')
            for c, h in enumerate(['Title', 'Price', 'URL'], 1):
                cell = ws3.cell(row=1, column=c, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align
            for r, p in enumerate(products, 2):
                ws3.cell(row=r, column=1, value=p.get('title','')).font = body_font
                ws3.cell(row=r, column=2, value=str(p.get('price',''))).font = body_font
                ws3.cell(row=r, column=3, value=p.get('url','')).font = body_font
            ws3.column_dimensions['A'].width = 40
            ws3.column_dimensions['B'].width = 12
            ws3.column_dimensions['C'].width = 50

            # Sheet4: Selling Points
            ws4 = wb.create_sheet('Selling Points')
            for c, h in enumerate(['#', 'Selling Point'], 1):
                cell = ws4.cell(row=1, column=c, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align
            for r, sp in enumerate(sp_list, 2):
                ws4.cell(row=r, column=1, value=r-1).font = body_font
                ws4.cell(row=r, column=2, value=sp.replace(chr(10), ' | ')).font = body_font
            ws4.column_dimensions['A'].width = 6
            ws4.column_dimensions['B'].width = 60

            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            return buf.read()

        with zipfile.ZipFile(zip_path, "w") as zf:
            # Word report
            zf.writestr("listing_report.docx", _make_docx())
            # Excel workbook
            zf.writestr("keywords_data.xlsx", _make_xlsx())
            # CSV products
            if products:
                csv_lines = ["title,price,url"]
                for p in products:
                    title = (p.get("title") or "").replace('"', '""')
                    price = str(p.get("price", "")).replace('"', '""')
                    url = (p.get("url") or "").replace('"', '""')
                    csv_lines.append(f'"{title}","{price}","{url}"')
                zf.writestr("products.csv", "\n".join(csv_lines))
            # Titles txt
            if titles:
                zf.writestr("titles.txt", "\n".join(titles))
            # Hashtags txt
            if hashtags:
                zf.writestr("hashtags.txt", "\n".join(hashtags))
            # Product photos
            for i in range(1, 10):
                img_path = os.path.join(CACHE_DIR, f"product_{i:02d}.jpg")
                if os.path.exists(img_path):
                    zf.write(img_path, f"product_{i:02d}.jpg")
            # Detail long image
            detail_path = os.path.join(CACHE_DIR, "detail_long.jpg")
            if os.path.exists(detail_path):
                zf.write(detail_path, "detail_long.jpg")

        return f"  ✅ 导出完成！包含：Word报告 + Excel数据 + 商品图 + 详情长图", zip_path
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"  导出失败：{e}", None


# ============================================================
# Build UI
# ============================================================

CSS = """
.gradio-container { max-width: 900px !important; margin: 0 auto !important; }
.status-text { font-size: 14px; color: #4A4A4A; }
"""

DROPDOWN_SLOTS_CONFIG = [
    ("dd_main_image", "产品主图"),
    ("dd_detail_1", "细节图 1"),
    ("dd_detail_2", "细节图 2"),
    ("dd_scene_image", "场景图"),
    ("dd_model_image", "模特图"),
]

with gr.Blocks(title="TK Listing Agent") as demo:
    # Shared session state
    session = gr.State({})

    gr.Markdown(
        "# TK Listing Agent — TikTok Shop 商品上架助手\n"
        "搜索商品、生成文案、制作主图 —— 一站式 TikTok Shop 上架工具。"
    )

    # ============================================================
    # Tab 1: 输入与搜索
    # ============================================================
    with gr.Tab("1. 输入与搜索") as tab1:
        with gr.Row():
            with gr.Column(scale=1):
                keyword = gr.Textbox(label="搜索关键词", placeholder="例如：yoga pants")
                country = gr.Dropdown(
                    choices=["US", "PH", "TH", "MY", "ID"], value="US", label="目标市场"
                )
                brand = gr.Textbox(label="品牌名称", placeholder="例如：ActivePro")
                category_kw = gr.Textbox(label="类目关键词", placeholder="例如：activewear")
            with gr.Column(scale=1):
                main_kw = gr.Textbox(label="主推关键词", placeholder="例如：yoga leggings")
                selling_pts = gr.Textbox(
                    label="核心卖点", placeholder="简要描述产品核心卖点",
                    lines=3,
                )
                material = gr.Textbox(label="材质", placeholder="例如：S925 纯银")
                series = gr.Textbox(label="系列名称", placeholder="例如：BIRTHSTONE SERIES")
                tagline = gr.Textbox(label="宣传语", placeholder="例如：Designed for everyday elegance.")
                material_desc = gr.Textbox(label="材质描述", placeholder="例如：Crafted with premium nylon blend for lasting comfort.")

        btn_search = gr.Button("🚀 开始搜索", variant="primary", size="lg")
        search_status = gr.Markdown("准备就绪。", elem_classes=["status-text"])

    # ============================================================
    # Tab 2: 商品列表
    # ============================================================
    with gr.Tab("2. 商品列表"):
        products_df = gr.Dataframe(
            headers=["商品标题", "价格", "链接"],
            label="Apify TikTok Shop 商品数据",
            interactive=False,
            wrap=True,
        )
        gr.Markdown("*数据来源：Apify TikTok Shop 搜索。返回 Tab 1 重新搜索。*")

    # ============================================================
    # Tab 3: 热帖趋势
    # ============================================================
    with gr.Tab("3. 热帖趋势"):
        posts_md = gr.Markdown("*暂无数据，请先在 Tab 1 中搜索。*")

    # ============================================================
    # Tab 4: 文案生成
    # ============================================================
    with gr.Tab("4. 文案生成"):
        crowd_kw = gr.Textbox(
            label="目标人群关键词",
            placeholder="例如：women, fitness lovers, yoga enthusiasts",
            value="women, fitness lovers",
        )
        btn_copy = gr.Button("🚀 生成文案", variant="primary", size="lg")
        copy_status = gr.Markdown("准备就绪。", elem_classes=["status-text"])

        gr.Markdown("### 产品标题")
        with gr.Row():
            title_seo = gr.Textbox(label="SEO 搜索标题")
            title_conv = gr.Textbox(label="转化标题")
            title_brand = gr.Textbox(label="品牌调性标题")

        gr.Markdown("### 卖点描述")
        desc_sp = gr.Markdown("*尚未生成。*")

        gr.Markdown("### 常见问答")
        desc_qa = gr.Markdown("*尚未生成。*")

        gr.Markdown("### 注意事项")
        desc_cautions = gr.Markdown("*尚未生成。*")

        gr.Markdown("### 话题标签")
        hashtags_out = gr.Textbox(label="话题标签", placeholder="#fyp #foryou...")

    # ============================================================
    # Tab 5: 产品图片
    # ============================================================
    with gr.Tab("5. 产品图片") as tab5:
        gr.Markdown("一键生成9张商品图（白底图、场景图、细节图、多角度图、模特图），基于参考图进行图生图。")
        ref_upload = gr.File(label="上传参考图片", file_types=["image"])
        product_keyword_input = gr.Textbox(
            label="产品名称（用于提示词描述）",
            placeholder="如：七脉轮手串、瑜伽裤",
        )
        with gr.Row():
            btn_product = gr.Button("🚀 生成9张商品图", variant="primary", size="lg")
            btn_clear_product = gr.Button("清除", variant="secondary")
        product_status = gr.Markdown("准备就绪。", elem_classes=["status-text"])
        product_gallery = gr.Gallery(label="生成的商品图", columns=3, rows=3, object_fit="contain")

    # ============================================================
    # Tab 6: 详情图生成
    # ============================================================
    with gr.Tab("6. 详情图") as tab6:
        gr.Markdown(
            "### 详情长图生成（13段完整详情页）\n"
            "**操作流程：**\n"
            "1. Tab 1 填写品牌、产品信息\n"
            "2. Tab 4 生成卖点文案\n"
            "3. Tab 5 上传参考图 → 生成9张商品图\n"
            "4. 回到 Tab 6，填写场景描述 → 点击生成详情长图"
        )
        with gr.Row():
            scene1 = gr.Textbox(label="场景1", placeholder="如：Meditation & Yoga", value="Daily Use")
            scene2 = gr.Textbox(label="场景2", placeholder="如：Daily Office Wear", value="Office Style")
            scene3 = gr.Textbox(label="场景3", placeholder="如：Weekend Casual", value="Weekend Casual")

        btn_detail = gr.Button("🤖 AI 生成详情长图（13段）", variant="primary", size="lg")
        detail_status = gr.Markdown("准备就绪。", elem_classes=["status-text"])
        detail_preview = gr.Image(label="详情长图预览", type="pil")

    # ============================================================
    # Tab 8: 导出打包
    # ============================================================
    with gr.Tab("8. 导出打包"):
        gr.Markdown("将已生成的所有内容（商品 CSV、标题、描述、话题标签、详情图）打包为 ZIP 文件。")
        btn_export = gr.Button("📦 导出全部为 ZIP", variant="secondary", size="lg")
        export_status = gr.Markdown("准备就绪。", elem_classes=["status-text"])
        export_file = gr.File(label="下载导出文件")

    # ============================================================
    # Wire callbacks
    # ============================================================

    btn_search.click(
        fn=on_search,
        inputs=[session, keyword, country],
        outputs=[session, search_status, products_df, posts_md],
    )

    btn_copy.click(
        fn=on_generate_copy,
        inputs=[session, brand, main_kw, category_kw, selling_pts, crowd_kw],
        outputs=[
            session, copy_status,
            title_seo, title_conv, title_brand,
            desc_sp, desc_qa, desc_cautions, hashtags_out,
        ],
    )

    # Tab5: generate product image → then refresh Tab7 dropdowns
    btn_product.click(
        fn=on_generate_product,
        inputs=[session, ref_upload, product_keyword_input, brand],
        outputs=[session, product_status, product_gallery],
    )

    btn_clear_product.click(
        fn=on_clear_product,
        inputs=[session],
        outputs=[session, product_status, product_gallery],
    )

    # Tab6: generate detail long image
    btn_detail.click(
        fn=on_generate_detail_v2,
        inputs=[session, brand, main_kw, material, tagline, material_desc,
                scene1, scene2, scene3],
        outputs=[session, detail_status, detail_preview],
    )

    btn_export.click(
        fn=on_export,
        inputs=[session],
        outputs=[export_status, export_file],
    )


# ============================================================
# Launch
# ============================================================

if __name__ == "__main__":
    demo.launch(
        server_name="0.0.0.0",
        server_port=int(os.getenv("PORT", "7861")),
        share=False,
        css=CSS,
    )
